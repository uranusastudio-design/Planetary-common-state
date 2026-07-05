from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\luus\Documents\Codex\2026-07-04\github\outputs\Revised_UCT_manuscript.docx")


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def add_p(doc, text="", style=None, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_font(run)
    if align is not None:
        p.alignment = align
    return p


def add_math(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, name="Cambria Math", size=11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_font(run)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_font(r, size=16 if level == 1 else 13 if level == 2 else 12, bold=True,
                 color="2E74B5" if level <= 2 else "1F4D78")
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    s = styles[name]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Unified Constraint Theory: An Effective Macroscopic Framework for Coupled Complex Systems")
set_font(r, size=18, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Revised manuscript for Physical Review E style")
set_font(r, size=11, italic=True, color="555555")

heading(doc, "Chapter 1", 1)
heading(doc, "Introduction", 2)
heading(doc, "1.1 Background", 3)
add_p(doc, "The long-term stability of complex systems is a central problem in nonequilibrium physics, Earth-system science, ecology, engineering, and network theory. These fields study different mechanisms, but they share a common concern: how large-scale coupled systems maintain macroscopic organization under persistent perturbations, finite resources, and irreversible dissipation.")
add_p(doc, "Established theories provide detailed descriptions of individual processes. Thermodynamics characterizes energy exchange and entropy production. Fluid mechanics describes momentum transport. Statistical mechanics connects microscopic dynamics to collective behavior. General relativity describes spacetime gravitation. Information theory quantifies uncertainty and communication. None of these theories is replaced or modified in the present work.")
add_p(doc, "The difficulty addressed here is representational. In many large-scale systems, observable degradation or loss of viability is not confined to one domain. It emerges from interactions among thermal, dynamical, chemical, structural, informational, energetic, geophysical, and viability-related constraints. A common effective representation is therefore useful when no single observable is sufficient.")
heading(doc, "1.2 Motivation", 3)
add_p(doc, "The absence of a unified representation is not caused by a lack of measurements. Radiative imbalance, circulation indices, chemical concentrations, biodiversity indicators, infrastructure metrics, and measures of adaptive capacity are all meaningful observables. However, they were developed for different scientific purposes and are often analyzed using domain-specific conventions.")
add_p(doc, "This fragmentation becomes especially important near possible regime transitions. Perturbations in one domain may change the sensitivity of another, and the relevant slow variables may be distributed across several observational channels. The present manuscript proposes a common macroscopic language for such coupled constraints without introducing new microscopic laws.")
heading(doc, "1.3 Scope of This Work", 3)
add_p(doc, "The objective of Unified Constraint Theory (UCT) is not to introduce a new physical force, a new fundamental law, or a replacement for thermodynamics, statistical mechanics, fluid mechanics, general relativity, quantum mechanics, information theory, or Earth-system science. UCT is proposed as an effective macroscopic framework for representing coupled constraints in large-scale complex systems.")
add_p(doc, "The framework is intended for systems whose behavior cannot be adequately summarized by one physical observable. Examples include Earth-system applications, ecological networks, infrastructure systems, and other socio-technical or environmental systems. Its scientific value depends on empirical calibration, falsifiable prediction, and comparison with simpler baselines.")
heading(doc, "1.4 Central Hypothesis", 3)
add_p(doc, "The central hypothesis is that heterogeneous observables can be interpreted as projections of a common effective constraint representation. The central object is the Unified Constraint Operator, denoted \\(\\mathbb{L}\\). It represents coupled macroscopic constraints, not a directly measured physical field.")
add_math(doc, "\\[ \\mathbb{L} \\]")
add_p(doc, "Observable quantities are projections of this operator. A thermal observable, a flow observable, or a viability indicator is not identified with \\(\\mathbb{L}\\) itself; rather, each is interpreted as one component extracted from the effective operator through a projection map.")
heading(doc, "1.5 Main Contributions", 3)
add_bullets(doc, [
    "A conservative effective-theory formulation of UCT centered on the Unified Constraint Operator \\(\\mathbb{L}\\).",
    "A projection framework in which observable constraint components are written as \\(L_i=\\pi_i(\\mathbb{L})\\).",
    "A constraint-coordinate geometry that represents cross-domain coupling without interpreting it as spacetime geometry.",
    "A data-calibrated implementation pathway using observation spaces, projection operators, assimilation, coupling estimation, hindcasting, and baseline comparison.",
    "A clear statement of limitations and falsifiability conditions."
])
heading(doc, "1.6 Organization of the Paper", 3)
add_p(doc, "Chapter 2 states the physical foundations and foundational principles. Chapter 3 defines the Unified Constraint Representation. Chapter 4 develops constraint geometry. Chapters 5-7 discuss effective dynamics, coupling, viability, and critical transitions. Chapters 8-10 describe implementation, assimilation, validation, and falsifiability. Chapter 11 states limitations and relations to established theories. Chapter 12 concludes. Appendices A-C collect the main formulas, normalization conventions, and representative data sources.")

heading(doc, "Chapter 2", 1)
heading(doc, "Physical Foundations and Foundational Principles of the Unified Constraint Theory", 2)
heading(doc, "2.1 Why Existing Physics Is Not Sufficient for This Representation", 3)
add_p(doc, "Existing physical theories are successful within their domains. The present claim is narrower: they do not by themselves provide a single effective coordinate system for comparing how heterogeneous constraints jointly affect macroscopic viability. This is a modeling gap, not a failure of established physics.")
add_p(doc, "UCT therefore asks a representational question: how can one describe the coupled constraint state of a large complex system using observables drawn from different physical and informational domains?")
heading(doc, "2.2 Principle I - Universality of Constraints", 3)
add_p(doc, "We propose that realistic macroscopic systems evolve under constraints arising from finite resources, dissipation, boundary conditions, structural limits, and interactions among subsystems. In this effective sense, the unconstrained idealization is not expected to describe Earth-scale or similarly complex systems.")
add_math(doc, "\\[ \\mathbb{L} \\neq 0 \\quad \\text{for realistic macroscopic systems.} \\]")
add_p(doc, "This statement is a modeling principle: constraints are treated as universal features of macroscopic evolution, not as new fundamental interactions.")
heading(doc, "2.3 Principle II - Constraints Are Not Physical Quantities", 3)
add_p(doc, "The operator \\(\\mathbb{L}\\) does not represent energy, entropy, mass, momentum, information, curvature of spacetime, or any conventional microscopic state variable. It is an effective mathematical object used to organize how these quantities constrain macroscopic evolution when they are coupled.")
add_p(doc, "Consequently, \\(\\mathbb{L}\\) is not directly measured. Only its observable projections are estimated from data.")
heading(doc, "2.4 Principle III - Observable Variables Are Projections", 3)
add_p(doc, "We define observable constraint components as projections of the Unified Constraint Operator:")
add_math(doc, "\\[ L_i = \\pi_i(\\mathbb{L}), \\qquad i=1,\\ldots,n. \\]")
add_p(doc, "The projection operator \\(\\pi_i\\) selects or constructs the component relevant to one observational domain. This interpretation allows conventional data products to remain conventional while also serving as components of a broader effective representation.")
heading(doc, "2.5 Principle IV - Coupling Generates Constraint Geometry", 3)
add_p(doc, "If projected components were statistically and dynamically independent, a Euclidean coordinate description would often be adequate. In coupled systems, however, perturbations in one projection can change the response of another. The geometry of the constraint-coordinate space is introduced to represent this coupling structure.")
add_math(doc, "\\[ ds_{\\mathbb{L}}^2 = g_{ij}(\\mathbf{L})\\,dL_i\\,dL_j. \\]")
add_p(doc, "Here \\(g_{ij}\\) is an effective constraint metric. It is inferred or modeled from cross-domain sensitivities; it is not a spacetime metric.")
heading(doc, "2.6 Principle V - Geometry Governs Macroscopic Evolution", 3)
add_p(doc, "We interpret macroscopic evolution as a trajectory in constraint-coordinate space. The trajectory is governed not only by changes in the projected components but also by the coupling structure represented by \\(g_{ij}\\).")
add_math(doc, "\\[ \\mathbf{L}(t) \\in \\mathcal{M}_{\\mathbb{L}}. \\]")
add_p(doc, "Transitions are therefore studied as changes in the estimated constraint state and its coupling geometry, rather than as evidence for a new physical law.")
heading(doc, "2.7 Scientific Position of the Theory", 3)
add_p(doc, "UCT is an effective, macroscopic, data-calibrated framework. It complements established physical theories by providing a common representation for heterogeneous constraints. Its validity depends on whether its projections, coupling estimates, and derived indices improve explanation or prediction relative to simpler alternatives.")

heading(doc, "Chapter 3", 1)
heading(doc, "Unified Constraint Representation", 2)
heading(doc, "3.1 Motivation", 3)
add_p(doc, "Large coupled systems are observed through many physical variables: temperature anomalies, energy imbalance, circulation indices, chemical concentrations, structural indicators, biodiversity measures, infrastructure data, and measures of adaptive capacity. These observables are meaningful, but they lack a common constraint representation.")
add_p(doc, "UCT addresses this by defining a shared effective object whose projections correspond to observable constraint components. The framework does not make heterogeneous observables identical; it places them in a common coordinate representation suitable for coupling analysis and empirical testing.")
heading(doc, "3.2 Unified Constraint Operator", 3)
add_p(doc, "Let \\(\\mathcal{S}\\) denote the system state space and \\(\\mathcal{C}\\) the effective constraint space. The Unified Constraint Operator is defined as")
add_math(doc, "\\[ \\mathbb{L}:\\mathcal{S}\\rightarrow\\mathcal{C}. \\]")
add_p(doc, "\\(\\mathbb{L}\\) is the central theoretical object of UCT. It is an effective operator representing coupled macroscopic constraints. It is not a second set of physical equations, not a new force, and not a replacement for the governing equations used within each domain.")
heading(doc, "3.3 Observable Constraint Projections", 3)
add_p(doc, "Observable constraint components are obtained by projection:")
add_math(doc, "\\[ L_i=\\pi_i(\\mathbb{L}), \\qquad i=1,\\ldots,n. \\]")
add_p(doc, "For the Earth-system application considered here, the projection vector is")
add_math(doc, "\\[ \\mathbf{L}=(L_T,L_F,L_C,L_S,L_I,L_E,L_G,L_V). \\]")
add_bullets(doc, [
    "\\(L_T\\): thermal constraint projection.",
    "\\(L_F\\): flow constraint projection.",
    "\\(L_C\\): chemical constraint projection.",
    "\\(L_S\\): structural constraint projection.",
    "\\(L_I\\): informational constraint projection.",
    "\\(L_E\\): energetic constraint projection.",
    "\\(L_G\\): gravitational or geophysical constraint projection.",
    "\\(L_V\\): viability constraint projection."
])
add_p(doc, "The viability projection is intended to represent macroscopic persistence and adaptive capacity, not a philosophical premise about civilization. Its empirical definition must be specified for each application.")
heading(doc, "3.4 Optional Scalar Constraint Index", 3)
add_p(doc, "A scalar aggregate may be useful for visualization, monitoring, or comparison, but it is not the fundamental object of the theory. If used, it is defined as a calibrated summary derived from the projection vector:")
add_math(doc, "\\[ L=F(\\mathbf{L}), \\qquad F:[0,1]^n\\rightarrow[0,1]. \\]")
add_p(doc, "The scalar \\(L\\) should be interpreted as an optional data-calibrated index. Its construction is application dependent and must be validated against independent observations or predictive tasks.")
add_p(doc, "In the linear baseline regime one may write")
add_math(doc, "\\[ L\\approx\\sum_i w_i L_i, \\qquad \\sum_i w_i=1. \\]")
add_p(doc, "Nonlinear forms of \\(F\\) may be considered only when supported by data and when they improve predictive performance relative to simpler baselines.")

heading(doc, "Chapter 4", 1)
heading(doc, "Constraint Geometry of Coupled Systems", 2)
heading(doc, "4.1 Constraint Manifold", 3)
add_p(doc, "The projection vector \\(\\mathbf{L}\\) may be treated as a coordinate representation of a constraint manifold, denoted \\(\\mathcal{M}_{\\mathbb{L}}\\). Each point on this manifold corresponds to an effective macroscopic constraint state, not a spatial location.")
add_math(doc, "\\[ \\mathbf{L}=(L_1,L_2,\\ldots,L_n)\\in\\mathcal{M}_{\\mathbb{L}}. \\]")
add_p(doc, "The manifold language is used to organize coupled coordinates and their sensitivities. It does not imply that UCT changes spacetime geometry.")
heading(doc, "4.2 Constraint Metric", 3)
add_p(doc, "A constraint metric may be introduced to quantify distances in constraint-coordinate space:")
add_math(doc, "\\[ ds_{\\mathbb{L}}^2=g_{ij}(\\mathbf{L})\\,dL_i\\,dL_j. \\]")
add_p(doc, "The metric \\(g_{ij}\\) represents the local weighting and coupling of projection changes. It must be estimated, modeled, or otherwise justified from empirical data. It is not a gravitational metric.")
heading(doc, "4.3 Constraint Distance and Constraint Velocity", 3)
add_p(doc, "A small constraint distance indicates two nearby macroscopic constraint states. A large constraint distance indicates substantial reorganization in the projected state space. The associated constraint velocity may be written")
add_math(doc, "\\[ V_{\\mathbb{L}}^2=g_{ij}(\\mathbf{L})\\frac{dL_i}{dt}\\frac{dL_j}{dt}. \\]")
add_p(doc, "This quantity measures the rate of motion through constraint-coordinate space. It is an effective diagnostic for comparing gradual evolution, rapid transition, or recovery.")
heading(doc, "4.4 Constraint Curvature", 3)
add_p(doc, "If the constraint manifold is treated as differentiable, one may define curvature quantities associated with \\(g_{ij}\\). These quantities summarize cross-domain coupling structure in constraint-coordinate space.")
add_p(doc, "Such curvature is not gravitational curvature. It does not modify Einstein's equation and should not be interpreted as a physical curvature of spacetime. Its use is justified only insofar as it improves the analysis of empirical coupling and macroscopic transition structure.")
heading(doc, "4.5 Constraint Potential", 3)
add_p(doc, "For some applications it may be useful to define an effective constraint potential \\(U(\\mathbf{L})\\) or \\(U(\\mathbb{L})\\). This potential is a modeling device for summarizing stability landscapes.")
add_math(doc, "\\[ \\nabla U(\\mathbf{L}_*)=0, \\qquad \\nabla^2 U(\\mathbf{L}_*)>0 \\]")
add_p(doc, "indicates a locally stable state in the chosen coordinate representation. The potential is not a microscopic energy function unless a separate derivation establishes such a relation.")

heading(doc, "Chapter 5", 1)
heading(doc, "Constraint Dynamics", 2)
heading(doc, "5.1 Evolution of the Constraint Operator", 3)
add_p(doc, "The Unified Constraint Operator represents an instantaneous effective constraint state. Its macroscopic evolution may be written as")
add_math(doc, "\\[ \\frac{D\\mathbb{L}}{Dt}=\\mathcal{F}(\\mathbb{L},\\Psi,t). \\]")
add_p(doc, "Here \\(\\Psi\\) denotes the collection of relevant physical state variables and boundary conditions, and \\(\\mathcal{F}\\) is an effective evolution operator to be inferred from data or constrained by domain models. It is not a microscopic law.")
heading(doc, "5.2 Constraint Sources", 3)
add_p(doc, "The effective operator \\(\\mathcal{F}\\) may receive contributions from thermal forcing, circulation, hydrology, chemistry, ecosystems, infrastructure, information flows, energy systems, and external boundary conditions. A schematic decomposition is")
add_math(doc, "\\[ \\mathcal{F}=\\sum_k \\mathcal{F}_k. \\]")
add_p(doc, "This decomposition is an accounting convention. The terms \\(\\mathcal{F}_k\\) must be defined through observables, models, or calibrated response functions.")
heading(doc, "5.3 Stable Evolution", 3)
add_p(doc, "Stable macroscopic evolution corresponds to slow change in the effective constraint state:")
add_math(doc, "\\[ \\left\\|\\frac{D\\mathbb{L}}{Dt}\\right\\|\\ll 1 \\]")
add_p(doc, "after specifying the norm and time scale. Rapid growth of this norm is a diagnostic candidate for instability, not by itself proof of an imminent transition.")
heading(doc, "5.4 Dynamic Equilibrium", 3)
add_p(doc, "An approximately stationary constraint state satisfies")
add_math(doc, "\\[ \\frac{D\\mathbb{L}}{Dt}\\approx 0. \\]")
add_p(doc, "This does not imply that physical processes cease. It indicates that the net projected constraint state is approximately stationary at the selected resolution.")

heading(doc, "Chapter 6", 1)
heading(doc, "Constraint Coupling Theory", 2)
heading(doc, "6.1 Coupling Matrix", 3)
add_p(doc, "The projected components are generally coupled. When differentiability is justified, the effective coupling matrix may be written")
add_math(doc, "\\[ C_{ij}=\\frac{\\partial \\pi_i(\\mathbb{L})}{\\partial \\pi_j(\\mathbb{L})}. \\]")
add_p(doc, "Equivalently, in projection coordinates, this can be represented locally as \\(C_{ij}=\\partial L_i/\\partial L_j\\). When differentiability is not justified, \\(C_{ij}\\) should be interpreted as an empirically estimated effective coupling coefficient.")
heading(doc, "6.2 Empirical Estimation", 3)
add_p(doc, "\\(C_{ij}(t)\\) may be estimated from lagged correlations, response functions, perturbation experiments, causal discovery methods, or data-assimilation sensitivity analysis. The estimation method should be reported together with uncertainty, lag structure, and robustness checks.")
heading(doc, "6.3 Physical Meaning", 3)
add_p(doc, "Examples include thermal-flow coupling, flow-structural coupling, structural-informational coupling, and energy-viability coupling. These coefficients quantify effective cross-domain interaction. They do not introduce a new interaction beyond the mechanisms already represented in the underlying domain models.")
heading(doc, "6.4 Symmetric and Asymmetric Coupling", 3)
add_p(doc, "In general, coupling need not be symmetric:")
add_math(doc, "\\[ C_{ij}\\neq C_{ji}. \\]")
add_p(doc, "For example, a thermal anomaly may strongly affect ecological or hydrological projections, while the reverse influence may occur only indirectly or at longer time scales.")
heading(doc, "6.5 Global Coupling Strength", 3)
add_p(doc, "A simple aggregate measure of off-diagonal coupling is")
add_math(doc, "\\[ \\Gamma=\\sum_{i\\neq j}|C_{ij}|. \\]")
add_p(doc, "The restriction \\(i\\neq j\\) avoids including diagonal self-coupling unless explicitly stated. Large \\(\\Gamma\\) indicates strong cross-domain coupling and may be associated with cascade risk, but the threshold for such risk must be empirically estimated.")

heading(doc, "Chapter 7", 1)
heading(doc, "Stability and Critical Transitions", 2)
heading(doc, "7.1 Viability Functional", 3)
add_p(doc, "Rather than defining stability solely as \\(1-L\\), UCT introduces a viability functional")
add_math(doc, "\\[ \\mathcal{V}=\\Phi(\\mathbb{L}) \\quad \\text{or} \\quad \\mathcal{V}=\\Phi(\\mathbf{L}). \\]")
add_p(doc, "\\(\\mathcal{V}\\) summarizes macroscopic persistence, adaptive capacity, or long-term stability as defined for the system under study. In the simplest linear approximation,")
add_math(doc, "\\[ \\mathcal{V}\\approx 1-L. \\]")
add_p(doc, "This approximation is useful only after \\(L\\) and its calibration have been specified.")
heading(doc, "7.2 Critical Threshold", 3)
add_p(doc, "A scalar threshold \\(L_c\\), if used, is empirical and system-dependent. It should be estimated from historical transitions, controlled simulations, or independent validation datasets.")
add_math(doc, "\\[ L<L_c \\quad \\text{stable regime}, \\qquad L\\approx L_c \\quad \\text{transition region}, \\qquad L>L_c \\quad \\text{post-threshold regime}. \\]")
add_p(doc, "These labels are diagnostic categories, not universal laws.")
heading(doc, "7.3 Constraint Potential and Critical Surfaces", 3)
add_p(doc, "In a multi-dimensional representation, critical behavior is more naturally associated with a surface or region in \\(\\mathbf{L}\\)-space than with one scalar number. A potential landscape, if used, should be calibrated to the data and tested against alternatives.")
add_math(doc, "\\[ \\mathcal{C}_{\\mathrm{crit}}=\\{\\mathbf{L}: \\Phi(\\mathbf{L})=\\Phi_c\\}. \\]")
heading(doc, "7.4 Early Warning Signals", 3)
add_p(doc, "If UCT captures relevant slow modes near a transition, recovery time should increase as the system approaches the estimated critical surface. Other empirical signals may include increasing variance, rising autocorrelation, changing cross-domain coupling, and reduced adaptive capacity.")
add_p(doc, "These are testable diagnostics. Their presence should be evaluated against null models and baseline indicators.")
heading(doc, "7.5 Cascading Transitions", 3)
add_p(doc, "A cascading transition occurs when perturbations propagate across projection domains through the coupling matrix. In this setting, a transition may be associated with increasing off-diagonal coupling,")
add_math(doc, "\\[ \\Gamma\\rightarrow\\Gamma_c, \\]")
add_p(doc, "where \\(\\Gamma_c\\) is an empirically estimated critical coupling level. No universal fixed value is implied.")

heading(doc, "Chapter 8", 1)
heading(doc, "Earth-System Implementation", 2)
heading(doc, "8.1 Observation Space", 3)
add_p(doc, "The Unified Constraint Operator cannot be measured directly. Its projections are estimated from heterogeneous observation spaces:")
add_math(doc, "\\[ \\Omega=\\{\\Omega_T,\\Omega_F,\\Omega_C,\\Omega_S,\\Omega_I,\\Omega_E,\\Omega_G,\\Omega_V\\}. \\]")
add_p(doc, "Each \\(\\Omega_i\\) denotes the observation space associated with one projection. Examples include satellite records, reanalysis products, in situ measurements, ecological inventories, and socioeconomic datasets.")
heading(doc, "8.2 Projection Operators", 3)
add_p(doc, "Implementation requires empirical projection operators")
add_math(doc, "\\[ L_i=\\mathcal{P}_i(\\Omega_i). \\]")
add_p(doc, "The operators \\(\\mathcal{P}_i\\) map observational data into dimensionless constraint projections. They may involve normalization, filtering, spatial aggregation, uncertainty propagation, or model-based feature extraction.")
add_bullets(doc, [
    "\\(L_T=\\mathcal{P}_T(Q_{\\mathrm{CERES}},T_{\\mathrm{anom}})\\) for thermal constraints.",
    "\\(L_F=\\mathcal{P}_F(U_{\\mathrm{ERA5}},\\omega,\\mathrm{AMOC})\\) for flow constraints.",
    "\\(L_C=\\mathcal{P}_C(\\mathrm{CO}_2,\\mathrm{CH}_4,\\mathrm{pH})\\) for chemical constraints.",
    "\\(L_S=\\mathcal{P}_S(\\mathrm{land},\\mathrm{ice},\\mathrm{infrastructure})\\) for structural constraints.",
    "\\(L_I=\\mathcal{P}_I(\\mathrm{biodiversity},\\mathrm{complexity})\\) for informational constraints.",
    "\\(L_V=\\mathcal{P}_V(\\mathrm{population},\\mathrm{adaptive\\ capacity})\\) for viability constraints."
])
heading(doc, "8.3 Unified State Estimation", 3)
add_p(doc, "The estimated operator is obtained through an assimilation operator:")
add_math(doc, "\\[ \\hat{\\mathbb{L}}=\\mathcal{H}(\\Omega). \\]")
add_p(doc, "\\(\\mathcal{H}\\) combines projected observations, uncertainty estimates, and coupling constraints to produce an empirical estimate of the macroscopic constraint state.")

heading(doc, "Chapter 9", 1)
heading(doc, "Data Assimilation Framework", 2)
heading(doc, "9.1 Multi-Source Assimilation", 3)
add_p(doc, "Observations differ in spatial resolution, temporal sampling, measurement uncertainty, and domain interpretation. A generalized assimilation formulation may therefore estimate \\(\\hat{\\mathbb{L}}\\) by minimizing an objective function:")
add_math(doc, "\\[ \\hat{\\mathbb{L}}=\\arg\\min_{\\mathbb{L}} J(\\mathbb{L}). \\]")
heading(doc, "9.2 Cost Function", 3)
add_p(doc, "A typical cost function may include a background term, an observation term, and a coupling-consistency term:")
add_math(doc, "\\[ J=J_b+J_o+J_c. \\]")
add_math(doc, "\\[ J_c=\\|C_{\\mathrm{obs}}-C_{\\mathrm{model}}\\|^2. \\]")
add_p(doc, "This term should be used only when the observed and modeled coupling matrices are defined consistently.")
heading(doc, "9.3 Uncertainty", 3)
add_p(doc, "Uncertainty in the estimated operator may be summarized by")
add_math(doc, "\\[ \\Sigma_{\\mathbb{L}}=\\mathrm{Cov}(\\hat{\\mathbb{L}}). \\]")
add_p(doc, "Uncertainty should be propagated through projections, scalar indices, coupling estimates, and hindcasting diagnostics.")
heading(doc, "9.4 Real-Time Updating", 3)
add_p(doc, "As new observations become available, the estimate may be updated by")
add_math(doc, "\\[ \\hat{\\mathbb{L}}_{k+1}=\\mathcal{U}(\\hat{\\mathbb{L}}_k,\\Omega_{k+1}). \\]")
add_p(doc, "\\(\\mathcal{U}\\) may be implemented using ensemble filters, variational assimilation, Bayesian updating, or other data-assimilation methods.")

heading(doc, "Chapter 10", 1)
heading(doc, "Numerical Framework and Validation", 2)
heading(doc, "10.1 Numerical Workflow", 3)
add_numbered(doc, [
    "Define observation spaces \\(\\Omega_i\\).",
    "Specify projection operators \\(\\mathcal{P}_i\\).",
    "Estimate projected components \\(L_i\\) and the operator \\(\\hat{\\mathbb{L}}\\).",
    "Estimate the coupling matrix \\(C_{ij}(t)\\).",
    "Construct or calibrate the constraint metric \\(g_{ij}(\\mathbf{L})\\).",
    "Evaluate prediction or diagnostic tasks.",
    "Validate against hindcasts and baseline models."
])
heading(doc, "10.2 Calibration", 3)
add_p(doc, "Projection operators, normalization functions, scalar aggregation, coupling coefficients, and metric components require empirical calibration. Calibration should be performed on training periods or training systems, with independent validation where possible.")
heading(doc, "10.3 Validation Metrics", 3)
add_p(doc, "Predictive performance may be evaluated using root-mean-square error, correlation, likelihood, Brier score, receiver-operating-characteristic measures, calibration curves, and temporal prediction skill. For continuous targets, a common metric is")
add_math(doc, "\\[ \\mathrm{RMSE}=\\sqrt{\\frac{1}{N}\\sum_{k=1}^N (y_k-\\hat{y}_k)^2}. \\]")
heading(doc, "10.4 Hindcasting", 3)
add_p(doc, "Hindcasting evaluates whether estimated constraint states and coupling measures recover known historical changes or regime shifts. A successful hindcast is not sufficient for validation, but failure to hindcast relevant events would challenge the usefulness of the framework.")
heading(doc, "10.5 Baseline Comparisons", 3)
add_p(doc, "UCT must be compared with simpler and established baselines. Relevant baselines include:")
add_bullets(doc, [
    "global mean surface temperature;",
    "top-of-atmosphere energy imbalance;",
    "AMOC indices;",
    "planetary-boundary indicators;",
    "PCA/EOF composite indices."
])
add_p(doc, "The framework is justified only if it offers interpretable or predictive advantages relative to such baselines.")
heading(doc, "10.6 Falsifiability", 3)
add_p(doc, "UCT is empirically testable. The framework would be challenged if no statistically significant mapping \\(\\Omega\\rightarrow\\hat{\\mathbb{L}}\\) can be established, if coupling estimates do not improve prediction or interpretation, if independent datasets yield inconsistent projections, or if simpler baselines perform as well or better.")

heading(doc, "Chapter 11", 1)
heading(doc, "Discussion and Limitations", 2)
heading(doc, "11.1 Scientific Position", 3)
add_p(doc, "UCT should be interpreted as an effective macroscopic representation. It complements existing theories by organizing heterogeneous constraints within a common projection space. It does not compete with microscopic physical laws or modify the equations used within established disciplines.")
heading(doc, "11.2 Relation to Existing Theories", 3)
add_p(doc, "Thermodynamics describes dissipation and energy exchange; statistical mechanics describes collective behavior arising from microscopic dynamics; fluid mechanics describes transport; general relativity describes gravitational spacetime; information theory describes uncertainty and communication; complex-systems theory studies feedbacks, networks, and transitions. UCT draws on this landscape but does not replace it.")
heading(doc, "11.3 Interpretation of the Constraint Operator", 3)
add_p(doc, "The Unified Constraint Operator \\(\\mathbb{L}\\) is not a physical field, conserved quantity, or fundamental interaction. It is an effective object whose observable projections are constructed from measurable variables. Its usefulness depends on whether this construction supports empirically testable analysis.")
heading(doc, "11.4 Scope of Applicability", 3)
add_p(doc, "The framework is most appropriate for macroscopic systems with heterogeneous interacting domains and sufficient observational coverage. It is less appropriate when a single well-established physical variable already provides a complete and predictive description.")
heading(doc, "11.5 Limitations", 3)
add_bullets(doc, [
    "UCT is presently a framework, not a validated predictive model.",
    "The choice of projections is not unique.",
    "The scalar index \\(L\\) may be calibration-dependent.",
    "The constraint metric \\(g_{ij}\\) requires empirical estimation.",
    "Coupling estimates may depend on time scale, lag structure, and data quality.",
    "The framework must outperform simpler baselines to justify its use."
])
heading(doc, "11.6 Future Work", 3)
add_p(doc, "Future work should focus on domain-specific projection definitions, uncertainty-aware data assimilation, hindcasting on independent historical periods, comparison with established indicators, and transparent release of calibration choices. The framework should be revised or rejected where empirical tests do not support it.")

heading(doc, "Chapter 12", 1)
heading(doc, "Conclusions", 2)
add_p(doc, "This manuscript has reformulated Unified Constraint Theory as a conservative effective framework for representing coupled constraints in macroscopic complex systems. The central object is the Unified Constraint Operator \\(\\mathbb{L}\\), whose observable quantities are projections \\(L_i=\\pi_i(\\mathbb{L})\\).")
add_p(doc, "The contribution is a unified effective representation, not a new fundamental law. UCT does not modify thermodynamics, statistical mechanics, fluid mechanics, general relativity, quantum mechanics, or information theory. Instead, it provides a data-calibrated language for organizing heterogeneous observables, estimating coupling, and testing whether a multi-domain representation improves understanding or prediction.")
add_p(doc, "The framework remains scientific only if its components are empirically specified and falsifiable. Its future value depends on transparent calibration, uncertainty quantification, hindcasting, and demonstrated performance relative to simpler baselines.")

heading(doc, "Appendix A", 1)
heading(doc, "Mathematical Foundations", 2)
heading(doc, "A.1 Unified Constraint Operator", 3)
add_math(doc, "\\[ \\mathbb{L}:\\mathcal{S}\\rightarrow\\mathcal{C}. \\]")
heading(doc, "A.2 Observable Projection", 3)
add_math(doc, "\\[ L_i=\\pi_i(\\mathbb{L}), \\qquad i=1,\\ldots,n. \\]")
heading(doc, "A.3 Projection Vector", 3)
add_math(doc, "\\[ \\mathbf{L}=(L_T,L_F,L_C,L_S,L_I,L_E,L_G,L_V). \\]")
heading(doc, "A.4 Optional Scalar Index", 3)
add_math(doc, "\\[ L=F(\\mathbf{L}), \\qquad F:[0,1]^n\\rightarrow[0,1]. \\]")
add_math(doc, "\\[ L\\approx\\sum_i w_i L_i, \\qquad \\sum_i w_i=1. \\]")
heading(doc, "A.5 Constraint Metric", 3)
add_math(doc, "\\[ ds_{\\mathbb{L}}^2=g_{ij}(\\mathbf{L})\\,dL_i\\,dL_j. \\]")
heading(doc, "A.6 Constraint Dynamics", 3)
add_math(doc, "\\[ \\frac{D\\mathbb{L}}{Dt}=\\mathcal{F}(\\mathbb{L},\\Psi,t). \\]")
heading(doc, "A.7 Coupling Matrix", 3)
add_math(doc, "\\[ C_{ij}=\\frac{\\partial \\pi_i(\\mathbb{L})}{\\partial \\pi_j(\\mathbb{L})}. \\]")
heading(doc, "A.8 Global Coupling", 3)
add_math(doc, "\\[ \\Gamma=\\sum_{i\\neq j}|C_{ij}|. \\]")
heading(doc, "A.9 Viability Functional", 3)
add_math(doc, "\\[ \\mathcal{V}=\\Phi(\\mathbb{L}) \\quad \\text{or} \\quad \\mathcal{V}=\\Phi(\\mathbf{L}). \\]")

heading(doc, "Appendix B", 1)
heading(doc, "Dimensionless Normalization", 2)
add_p(doc, "Different observables possess different physical units and uncertainty structures. UCT therefore uses dimensionless projected components, typically scaled to \\([0,1]\\), while recognizing that normalization is domain-specific.")
add_math(doc, "\\[ L_i=\\frac{x_i-x_i^{\\mathrm{ref}}}{x_i^{\\mathrm{crit}}-x_i^{\\mathrm{ref}}}. \\]")
add_p(doc, "When a bounded index is required, clipping may be applied:")
add_math(doc, "\\[ L_i\\leftarrow \\min(1,\\max(0,L_i)). \\]")
add_p(doc, "Alternatively, a monotonic transform may be used:")
add_math(doc, "\\[ L_i=\\sigma_i(x_i). \\]")
add_p(doc, "The reference value \\(x_i^{\\mathrm{ref}}\\), critical value \\(x_i^{\\mathrm{crit}}\\), and transform \\(\\sigma_i\\) must be chosen and reported for each projection domain. Spatial scaling, temporal scaling, energy scaling, and information scaling can all be incorporated through projection operators.")

heading(doc, "Appendix C", 1)
heading(doc, "Representative Observational Data Sources", 2)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
hdr = table.rows[0].cells
for i, text in enumerate(["Projection", "Observable", "Example datasets"]):
    hdr[i].text = text
    set_cell_shading(hdr[i], "F2F4F7")
    set_cell_margins(hdr[i])
    hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in hdr[i].paragraphs:
        for r in p.runs:
            set_font(r, bold=True)

rows = [
    ("Thermal", "TOA imbalance, temperature anomaly", "CERES, ERA5, NOAA"),
    ("Flow", "wind, vorticity, AMOC proxies", "ERA5, ECMWF, Argo, RAPID"),
    ("Chemical", "CO2, CH4, ocean pH", "NOAA, Mauna Loa, SOCAT"),
    ("Structural", "land cover, ice mass, infrastructure", "MODIS, Landsat, Sentinel, GRACE, ICESat"),
    ("Informational", "biodiversity, ecosystem complexity", "GBIF, IUCN, Living Planet Database"),
    ("Energetic", "energy availability, energy efficiency", "IEA, national energy statistics"),
    ("Geophysical", "geophysical boundary conditions", "GRACE, geodesy, lithosphere datasets"),
    ("Viability", "population, infrastructure, adaptive capacity", "UN, World Bank, national statistics"),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cells[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_font(r)

widths = [Inches(1.35), Inches(2.55), Inches(2.6)]
for row in table.rows:
    for cell, width in zip(row.cells, widths):
        cell.width = width
set_table_geometry(table, [1944, 3672, 3744], indent_dxa=120)

add_p(doc, "These examples are illustrative. A validated application must specify data provenance, preprocessing, uncertainty treatment, and the projection operator used for each domain.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
